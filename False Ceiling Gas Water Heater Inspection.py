import streamlit as st
import os
import zipfile
import shutil
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from io import BytesIO

# --- 1. Helper Functions ---

def create_embedded_template(save_path):
    doc = Document()
    heading = doc.add_heading('Towngas Inspection Record', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    labels = ["Project Name/Location", "Flat", "Name of Inspector", "Inspection Date"]
    
    for i, label in enumerate(labels):
        cell = table.cell(i, 0)
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.save(save_path)

def parse_filename_with_zeros(filename):
    """
    Parses filenames using the '0' rule.
    Format: Project-Tower-Flat-Inspector-Date
    Rule: If Tower or Flat is '0', treat it as empty.
    """
    base_name = os.path.splitext(filename)[0]
    # Remove counters like " (2)"
    clean_name = re.sub(r'\s\(\d+\)$', '', base_name)
    parts = clean_name.split('-')
    
    # Initialize defaults
    project = ""
    tower = ""
    flat = ""
    inspector = ""
    date = ""
    
    # Logic: We expect at least 5 parts (Project, Tower, Flat, Inspector, Date)
    if len(parts) >= 5:
        project = parts[0]
        
        # --- THE 0 RULE ---
        # If part is '0', set to empty string, otherwise keep the text
        tower_raw = parts[1]
        flat_raw = parts[2]
        
        tower = "" if tower_raw == '0' else tower_raw
        flat = "" if flat_raw == '0' else flat_raw
        
        inspector = parts[3]
        date = '-'.join(parts[4:]) # Join remaining parts for date
    else:
        # Fallback for filenames that don't match the new standard
        # We try to grab the first part as project at minimum
        if len(parts) > 0: project = parts[0]

    return {
        "filename": filename,
        "Project": project,
        "Tower": tower,
        "Flat": flat,
        "Inspector": inspector,
        "Date": date,
        "full_path": "" # To be filled during file walk
    }

# --- 2. Main Streamlit App ---

def main():
    st.set_page_config(page_title="Inspection Report Generator", page_icon="📝", layout="wide")
    
    st.title("📝 Gas Water Heater Inspection Record Generator")
    st.markdown("""
    **Instructions:**
    1. Rename your images using this standard format:  
       `Project-Tower-Flat-Inspector-Date.jpg`
    2. **The "0" Rule:**
       * If there is no Tower, put `0`.
       * If there is no Flat, put `0`.
    
    **Examples:**
    * Standard: `太湖花園-5座-1A-譚大文-20-01-2025`
    * No Tower: `NKIL-0-1A-陳明-20-01-2025` (Tower becomes empty)
    * No Tower & No Flat: `147 Waterloo Road-0-0-陳明-20-01-2025` (Both empty)
    """)

    # Initialize Session State
    if 'processed_data' not in st.session_state:
        st.session_state.processed_data = None
    if 'temp_dir_obj' not in st.session_state:
        st.session_state.temp_dir_obj = None

    # Step 1: Upload
    uploaded_file = st.file_uploader("1. Upload Images ZIP", type="zip")

    if uploaded_file:
        # Check if we need to process this new file
        if st.button("2. Process & Review Images", type="primary"):
            
            # Create persistent temp directory for this session
            st.session_state.temp_dir_obj = tempfile.TemporaryDirectory()
            temp_dir = st.session_state.temp_dir_obj.name
            
            extract_path = os.path.join(temp_dir, "input_images")
            os.makedirs(extract_path, exist_ok=True)
            
            # Extract
            with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
                zip_ref.extractall(extract_path)
                
            # Parse Files
            parsed_records = []
            valid_extensions = ('.png', '.jpg', '.jpeg')
            
            for root_dir, dirs, files in os.walk(extract_path):
                for file in files:
                    if file.lower().endswith(valid_extensions):
                        # Use the new "0" logic function
                        record = parse_filename_with_zeros(file)
                        record['full_path'] = os.path.join(root_dir, file)
                        parsed_records.append(record)
            
            if not parsed_records:
                st.error("No valid images found in ZIP.")
            else:
                # Store in session state as DataFrame
                st.session_state.processed_data = pd.DataFrame(parsed_records)
                st.success(f"Processed {len(parsed_records)} images.")

    # Step 2: Review & Edit
    if st.session_state.processed_data is not None:
        st.divider()
        st.subheader("3. Review & Edit Details")
        st.info("👇 Check the table below. '0' inputs should now be empty cells.")
        
        # Display Editable Table
        edited_df = st.data_editor(
            st.session_state.processed_data,
            column_order=("Project", "Tower", "Flat", "Inspector", "Date"),
            disabled=["filename"], 
            num_rows="dynamic",
            use_container_width=True,
            hide_index=True
        )

        # Step 3: Generate
        st.divider()
        if st.button("4. Confirm & Generate Reports", type="primary"):
            
            temp_dir = st.session_state.temp_dir_obj.name
            output_path = os.path.join(temp_dir, "output_docs")
            os.makedirs(output_path, exist_ok=True)
            
            # Create Template
            template_path = os.path.join(temp_dir, "template.docx")
            create_embedded_template(template_path)
            
            # Group by Unique Location
            grouped = {}
            
            for index, row in edited_df.iterrows():
                p = row['Project'].strip()
                t = row['Tower'].strip() if row['Tower'] else ""
                f = row['Flat'].strip() if row['Flat'] else ""
                
                # Create a unique key. 
                # Note: This key is internal logic only. 
                # The filename will be generated separately later.
                key = f"{p}-{t}-{f}"
                
                if key not in grouped:
                    grouped[key] = {
                        'project': p,
                        'tower': t,
                        'flat': f,
                        'inspector': row['Inspector'],
                        'date': row['Date'],
                        'images': []
                    }
                
                img_path = st.session_state.processed_data.loc[index, 'full_path']
                grouped[key]['images'].append(img_path)

            # Generate Documents
            progress_bar = st.progress(0)
            total_groups = len(grouped)
            
            zip_buffer = BytesIO()
            
            try:
                for i, (key, data) in enumerate(grouped.items()):
                    doc = Document(template_path)
                    table = doc.tables[0]
                    
                    table.cell(0, 1).text = data['project']
                    
                    # --- Table Logic: "Tower Flat" ---
                    # e.g. "5座 1A", or "1A" (if tower empty), or "" (if both empty)
                    location_parts = []
                    if data['tower']: location_parts.append(data['tower'])
                    if data['flat']: location_parts.append(data['flat'])
                    
                    table.cell(1, 1).text = " ".join(location_parts)
                    
                    table.cell(2, 1).text = str(data['inspector'])
                    table.cell(3, 1).text = str(data['date'])
                    
                    # Add Images
                    data['images'].sort()
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.2
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    
                    for img_path in data['images']:
                        try:
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(2.5))
                            run.add_text(" " * 8)
                        except Exception as e:
                            st.warning(f"Skipped image in {key}: {e}")
                            
                    # --- Filename Logic ---
                    # Join non-empty parts with hyphens
                    # e.g. ["NKIL", "1A"] -> "NKIL-1A.docx"
                    # e.g. ["147 Waterloo"] -> "147 Waterloo.docx"
                    filename_parts = [data['project']]
                    if data['tower']: filename_parts.append(data['tower'])
                    if data['flat']: filename_parts.append(data['flat'])
                    
                    safe_filename_base = "-".join(filename_parts)
                    safe_filename = f"{safe_filename_base}.docx".replace('/', '_')
                    
                    doc.save(os.path.join(output_path, safe_filename))
                    progress_bar.progress((i + 1) / total_groups)
                
                # Zip Creation
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, _, files in os.walk(output_path):
                        for file in files:
                            zf.write(os.path.join(root, file), arcname=file)
                
                zip_buffer.seek(0)
                st.success("✅ Reports Generated Successfully!")
                
                st.download_button(
                    label="⬇️ Download Reports (ZIP)",
                    data=zip_buffer,
                    file_name="Inspection_Reports.zip",
                    mime="application/zip"
                )
                
            except Exception as e:
                st.error(f"Error during generation: {e}")

if __name__ == "__main__":
    main()



